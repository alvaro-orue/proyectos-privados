"""
Utilidad para convertir archivos Markdown a Word (.docx)
Convierte automáticamente los informes generados a formato Word
"""
import os
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("⚠️  Librería 'python-docx' no instalada")
    print("📦 Instalar con: pip install python-docx")
    sys.exit(1)


def markdown_to_word(markdown_file, output_file=None):
    """
    Convierte un archivo Markdown a Word (.docx)

    Args:
        markdown_file (str): Ruta al archivo .md
        output_file (str): Ruta al archivo .docx de salida (opcional)

    Returns:
        str: Ruta al archivo Word generado
    """
    if not os.path.exists(markdown_file):
        raise FileNotFoundError(f"Archivo no encontrado: {markdown_file}")

    # Determinar nombre del archivo de salida
    if output_file is None:
        output_file = markdown_file.replace('.md', '.docx')

    # Leer el archivo Markdown
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()

    # Crear documento Word
    doc = Document()

    # Configurar estilos
    configure_styles(doc)

    # Procesar el contenido Markdown línea por línea
    lines = markdown_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Título principal (# )
        if line.startswith('# '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Título nivel 2 (## )
        elif line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)

        # Título nivel 3 (### )
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)

        # Título nivel 4 (#### )
        elif line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=4)

        # Lista con viñetas (- o *)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Limpiar emojis y formato
            text = clean_markdown_formatting(text)
            doc.add_paragraph(text, style='List Bullet')

        # Lista numerada (1. )
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = clean_markdown_formatting(text)
            doc.add_paragraph(text, style='List Number')

        # Bloque de código (```)
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1

            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'Quote'

            # Formato de código
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 128)

        # Tabla Markdown (| )
        elif line.strip().startswith('|') and '|' in line:
            # Detectar inicio de tabla
            table_lines = [line]
            i += 1

            # Capturar todas las líneas de la tabla
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1  # Retroceder una línea

            # Crear tabla en Word
            create_table_from_markdown(doc, table_lines)

        # Línea horizontal (---)
        elif line.strip() == '---' or line.strip() == '___':
            doc.add_paragraph('_' * 80)

        # Texto normal
        elif line.strip():
            text = clean_markdown_formatting(line)
            if text:
                p = doc.add_paragraph(text)
                # Detectar si es una línea de estado/resultado
                if '✅' in text or '❌' in text or '⚠️' in text:
                    for run in p.runs:
                        run.font.bold = True

        # Línea vacía
        else:
            doc.add_paragraph()

        i += 1

    # Guardar documento
    doc.save(output_file)

    return output_file


def configure_styles(doc):
    """Configurar estilos del documento"""
    # Estilo para títulos
    styles = doc.styles

    # Modificar estilo de Heading 1
    if 'Heading 1' in styles:
        style = styles['Heading 1']
        style.font.size = Pt(18)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)

    # Modificar estilo de Heading 2
    if 'Heading 2' in styles:
        style = styles['Heading 2']
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 102, 204)


def clean_markdown_formatting(text):
    """Limpiar formato Markdown del texto"""
    # Remover negrita (**text**)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)

    # Remover cursiva (*text* o _text_)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)

    # Remover código inline (`text`)
    text = re.sub(r'`(.+?)`', r'\1', text)

    # Remover enlaces [text](url)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)

    return text


def create_table_from_markdown(doc, table_lines):
    """
    Crear tabla en Word desde líneas Markdown

    Args:
        doc: Documento Word
        table_lines: Lista de líneas de la tabla Markdown
    """
    # Filtrar líneas de separador (|---|---|)
    table_lines = [line for line in table_lines if not re.match(r'^\|[\s\-:]+\|', line)]

    if not table_lines:
        return

    # Parsear filas
    rows_data = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Ignorar primero y último vacío
        rows_data.append(cells)

    if not rows_data:
        return

    # Crear tabla
    num_rows = len(rows_data)
    num_cols = len(rows_data[0])

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Llenar tabla
    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = clean_markdown_formatting(cell_text)

            # Negrita para encabezados (primera fila)
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

                # Color de fondo para encabezados
                cell._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
                )


def parse_xml(xml_string):
    """Helper para parsear XML"""
    from docx.oxml import parse_xml as docx_parse_xml
    return docx_parse_xml(xml_string)


def nsdecls(*prefixes):
    """Helper para namespace declarations"""
    from docx.oxml.ns import nsdecls as docx_nsdecls
    return docx_nsdecls(*prefixes)


def convert_all_markdown_in_directory(directory='.', recursive=False):
    """
    Convierte todos los archivos .md en un directorio a .docx

    Args:
        directory (str): Directorio a buscar
        recursive (bool): Buscar en subdirectorios

    Returns:
        list: Lista de archivos Word generados
    """
    converted_files = []

    pattern = '**/*.md' if recursive else '*.md'

    for md_file in Path(directory).glob(pattern):
        if md_file.name.startswith('.'):  # Ignorar archivos ocultos
            continue

        try:
            print(f"📄 Convirtiendo: {md_file.name}")
            docx_file = markdown_to_word(str(md_file))
            converted_files.append(docx_file)
            print(f"✅ Generado: {Path(docx_file).name}")
        except Exception as e:
            print(f"❌ Error convirtiendo {md_file.name}: {e}")

    return converted_files


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description='Convertir Markdown a Word')
    parser.add_argument('input', nargs='?', help='Archivo .md de entrada')
    parser.add_argument('-o', '--output', help='Archivo .docx de salida')
    parser.add_argument('-d', '--directory', help='Convertir todos los .md en un directorio')
    parser.add_argument('-r', '--recursive', action='store_true', help='Buscar recursivamente')

    args = parser.parse_args()

    if args.directory:
        # Convertir todos los archivos en directorio
        print(f"🔍 Buscando archivos .md en: {args.directory}")
        converted = convert_all_markdown_in_directory(args.directory, args.recursive)
        print(f"\n✅ {len(converted)} archivo(s) convertido(s)")

    elif args.input:
        # Convertir un solo archivo
        try:
            output = markdown_to_word(args.input, args.output)
            print(f"✅ Archivo Word generado: {output}")
        except Exception as e:
            print(f"❌ Error: {e}")
            sys.exit(1)

    else:
        print("⚠️  Uso:")
        print("  python markdown_to_word.py archivo.md")
        print("  python markdown_to_word.py -d SCRIPTS_PARA_QA")
        print("  python markdown_to_word.py archivo.md -o salida.docx")
